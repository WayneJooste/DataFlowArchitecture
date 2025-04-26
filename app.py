import streamlit as st
from components.architecture_diagram import display_architecture_diagram
from components.component_descriptions import display_component_descriptions
from components.data_sources import display_data_sources
from components.reports_dashboards import display_reports_dashboards
from components.superset_intro import display_superset_intro
from components.potv_ce import display_potv_ce
from docx import Document
from docx.shared import Inches

# Page configuration

st.set_page_config(
    page_title="BI Architecture Migration: DOMO to Amazon Web Services",
    page_icon="assets/favicon.png",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Sidebar for navigation
st.sidebar.title("Navigation")
page = st.sidebar.radio(
    "Select a section:",
    [
        "Architecture Overview", 
        "Component Descriptions", 
        "Data Sources", 
        "Standard Reports & Dashboards",
        "POTV - Superset intro",
        "Architecture Cost Estimates"
    ]
)
# Word File Creation Function
def create_word_file():
    doc = Document()
    doc.add_picture('assets/logo.jpg', width=Inches(1.5))  # or logo.png if you want
    doc.add_heading('BI Architecture Migration Tool', 0)
    doc.add_paragraph('Created by Data Alchemist')
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph('This tool assists with migrating BI systems from DOMO to AWS.')
    doc.add_heading('Architecture Overview', level=1)
    doc.add_paragraph('This diagram illustrates the migration from DOMO to a modern AWS-based BI architecture.')
    doc.add_heading('Component Descriptions', level=1)
    doc.add_paragraph('Detailed descriptions of each technology component.')
    doc.add_heading('Data Sources', level=1)
    doc.add_paragraph('Overview of the various data sources and integration into AWS architecture.')
    doc.add_heading('Standard Reports & Dashboards', level=1)
    doc.add_paragraph('Recommended standard reports and dashboards for an online vaporization business.')
    doc.add_heading('Architecture Cost Estimates', level=1)
    doc.add_paragraph('Estimated costs for architecture build.')

    doc.save('output.docx')

# --- Add Word Download Button Below Navigation ---
st.sidebar.markdown("---")  # horizontal line for separation
st.sidebar.subheader("Export Options")
if st.sidebar.button('Generate Word Report'):
    # Call your Word file creation function
    create_word_file()
    with open('output.docx', 'rb') as f:
        st.sidebar.download_button('Download Report', f, file_name='BI_Migration_Plan.docx')


# Page title and intro
st.title("BI Architecture Migration: DOMO to Amazon Web Services (AWS)")

# Display the selected page
if page == "Architecture Overview":
    st.header("Architecture Diagram")
    st.write("""
    This diagram illustrates the proposed migration from DOMO to a modern AWS-based BI architecture 
    using AlcheSanctum, AlcheFlow, and Apache Superset. It shows both the current phase and 
    future expansion to incorporate AI/ML and real-time analytics.
    """)
    display_architecture_diagram()

elif page == "Component Descriptions":
    st.header("Technology Components")
    st.write("""
    Detailed descriptions of each technology component in the proposed architecture.
    """)
    display_component_descriptions()

elif page == "Data Sources":
    st.header("Data Sources & Integration")
    st.write("""
    Overview of the various data sources and how they'll be integrated into the new architecture.
    """)
    display_data_sources()

elif page == "Standard Reports & Dashboards":
    st.header("Standard Reports & Dashboards")
    st.write("""
    Recommended standard reports and dashboards for an online vaporization business.
    """)
    display_reports_dashboards()

elif page == "POTV - Superset intro":
    st.header("POTV - Superset intro")
    st.write("""
    An introduction video to Superset.
    """)
    display_superset_intro()

elif page == "Architecture Cost Estimates":
    st.header("Architecture Cost Estimates")
    st.write("""
    Estimated costs for Architecture build.
    """)
    display_potv_ce()

# Footer
st.markdown("---")
st.image("assets/logo.jpg", width=150)
st.markdown(
    """
    <div style="text-align: left;">
        BI Architecture Migration Planning Tool | Created by Data Alchemist
    </div>
    """,
    unsafe_allow_html=True
)

